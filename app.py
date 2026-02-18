import streamlit as st
import pdfplumber
import google.generativeai as genai
import json
import re
import html
import tempfile
import os
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Cm, Pt, RGBColor
import io
from datetime import datetime, timedelta
import pandas as pd
from streamlit_option_menu import option_menu

# --- CONFIGURAÇÃO INICIAL ---
st.set_page_config(page_title="PeritoSaaS Pro", page_icon="⚖️", layout="wide")

# --- CONSTANTES ---
GEMINI_MODEL = "models/gemini-flash-latest"

# --- CSS E LAYOUT ---
st.markdown("""
<style>
    .block-container { padding-top: 4rem !important; padding-bottom: 5rem; }
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    .stButton > button { width: 100%; border-radius: 5px; }
</style>
""", unsafe_allow_html=True)


# --- FUNÇÕES DE TEMPLATES (GERADORES) ---
def criar_template_aceite():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(3)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2)
    doc.add_heading('PETIÇÃO DE ACEITE', 0)
    doc.add_paragraph('Excelentíssimo Senhor Doutor Juiz de Direito da {{vara}}')
    doc.add_paragraph('\nProcesso nº: {{numero_processo}}')
    doc.add_paragraph('Autor: {{autor}}')
    doc.add_paragraph('Réu: {{reu}}')
    doc.add_paragraph('\n{{ nome_perito }}, perito nomeado nos autos em epígrafe, vem, respeitosamente, perante Vossa Excelência, ACEITAR o honroso encargo para o qual foi designado.')
    doc.add_paragraph('\nRequer a juntada de seus dados bancários e contatos profissionais em anexo.')
    doc.add_paragraph('\nNestes termos,\nPede deferimento.')
    doc.add_paragraph('\nBelém, {{ data_atual }}.')
    doc.add_paragraph('\n___________________________\n{{ nome_perito }}\nPerito do Juízo')
    return doc


def criar_template_honorarios():
    doc = Document()
    doc.add_heading('PROPOSTA DE HONORÁRIOS', 0)
    doc.add_paragraph('Excelentíssimo Juiz da {{vara}}')
    doc.add_paragraph('Processo: {{numero_processo}}')
    doc.add_paragraph('\nO Perito vem apresentar sua estimativa de honorários baseada na complexidade do trabalho:')
    doc.add_paragraph('\n1. Vistoria Técnica: {{horas_vistoria}} horas')
    doc.add_paragraph('2. Análise Documental: {{horas_analise}} horas')
    doc.add_paragraph('3. Redação do Laudo: {{horas_redacao}} horas')
    doc.add_paragraph('TOTAL DE HORAS ESTIMADAS: {{total_horas}}h')
    doc.add_paragraph('\nValor da Hora Técnica: R$ {{valor_hora}}')
    doc.add_paragraph('VALOR TOTAL DOS HONORÁRIOS: R$ {{valor_total}}')
    doc.add_paragraph('\nNestes termos,\nPede deferimento.')
    doc.add_paragraph('\n{{ nome_perito }}')
    return doc


# --- FUNÇÃO DE DATA ---
def calcular_prazo_uteis(data_inicio, dias):
    """Calcula prazo em dias úteis (sem feriados — considere usar workalendar para maior precisão)."""
    dias_uteis = 0
    data_atual = data_inicio
    while dias_uteis < dias:
        data_atual += timedelta(days=1)
        if data_atual.weekday() < 5:
            dias_uteis += 1
    return data_atual


# --- FUNÇÃO: SALVAR DOC E ABRIR COM DOCXTPL VIA ARQUIVO TEMPORÁRIO ---
def renderizar_docx(doc, ctx):
    """Salva o Document em arquivo temporário e renderiza com DocxTemplate de forma segura."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        doc.save(tmp_path)

    doc_tpl = DocxTemplate(tmp_path)
    doc_tpl.render(ctx)

    bio_final = io.BytesIO()
    doc_tpl.save(bio_final)

    os.unlink(tmp_path)  # Remove arquivo temporário

    return bio_final


# --- SETUP API (robusto) ---
api_key = os.environ.get("GEMINI_API_KEY")
if not api_key:
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
    except Exception:
        api_key = None

if not api_key:
    st.error("⚠️ API Key não configurada. Defina GEMINI_API_KEY nas variáveis de ambiente ou em st.secrets.")
    st.stop()

# Configura a API uma única vez na inicialização
genai.configure(api_key=api_key)


# --- NOME DO PERITO (configurável) ---
if "nome_perito" not in st.session_state:
    st.session_state.nome_perito = ""

with st.sidebar:
    st.markdown("### ⚙️ Configurações")
    st.session_state.nome_perito = st.text_input(
        "Seu nome completo (Perito)",
        value=st.session_state.nome_perito,
        placeholder="Ex: Dr. João Silva"
    )
    if not st.session_state.nome_perito:
        st.warning("Informe seu nome para gerar documentos.")


# --- MENU ---
selected = option_menu(
    menu_title=None,
    options=["Dashboard de Processos", "Ferramentas Manuais"],
    icons=["kanban", "tools"],
    default_index=0,
    orientation="horizontal",
    styles={
        "container": {"padding": "5px", "background-color": "#262730"},
        "icon": {"color": "#ffffff", "font-size": "20px"},
        "nav-link": {"font-size": "16px", "text-align": "center", "margin": "0px", "color": "#ffffff"},
        "nav-link-selected": {"background-color": "#4e91d6"},
    }
)


# ==============================================================================
# DASHBOARD
# ==============================================================================
if selected == "Dashboard de Processos":
    st.markdown("### 🗂️ Análise de Autos")
    st.markdown("O sistema identificará **Nomeações**, **Quesitos** e **Intimações** e oferecerá a ferramenta certa.")

    uploaded_file_integral = st.file_uploader("📂 Suba o PDF Completo", type="pdf", key="pdf_integral")

    if uploaded_file_integral and st.button("🔍 Analisar Autos", type="primary"):
        with st.spinner("Lendo o processo inteiro (isso pode levar um minuto)..."):
            try:
                model = genai.GenerativeModel(GEMINI_MODEL)

                with pdfplumber.open(uploaded_file_integral) as pdf:
                    texto_paginado = ""
                    for i, page in enumerate(pdf.pages):
                        txt = page.extract_text()
                        if txt:
                            texto_paginado += f"--- PÁGINA {i+1} ---\n{txt}\n"

                prompt = f"""
                Atue como Assistente Pericial. Analise o processo e identifique eventos chave.

                1. DADOS DO PROCESSO: Extraia Numero, Autor, Réu e Vara.
                2. EVENTOS:
                   - NOMEACAO: Se houve nomeação.
                   - QUESITOS: Se há perguntas das partes (copie-as ÍPSIS LITTERIS).
                   - INTIMACAO: Se há prazo para proposta de honorários ou laudo.

                Retorne APENAS JSON válido, sem texto antes ou depois, sem blocos de código:
                {{
                    "metadados": {{ "numero": "...", "autor": "...", "reu": "...", "vara": "..." }},
                    "tarefas": [
                        {{
                            "tipo": "NOMEACAO",
                            "pagina": "45",
                            "descricao": "Nomeado para perícia médica."
                        }},
                        {{
                            "tipo": "QUESITOS",
                            "pagina": "52",
                            "descricao": "Quesitos do Autor",
                            "lista_quesitos": ["1. O periciando...", "2. Há nexo..."]
                        }},
                        {{
                            "tipo": "HONORARIOS",
                            "pagina": "60",
                            "descricao": "Intimado para apresentar proposta."
                        }}
                    ]
                }}
                TEXTO: {texto_paginado}
                """

                resp = model.generate_content(prompt)

                # Parser robusto: extrai bloco JSON mesmo que venha com texto ao redor
                texto_limpo = resp.text.strip()
                match = re.search(r'\{.*\}', texto_limpo, re.DOTALL)
                if not match:
                    st.error("A IA retornou uma resposta inesperada. Tente novamente.")
                    st.stop()

                st.session_state.dashboard_dados = json.loads(match.group())
                st.success("Análise concluída!")

            except json.JSONDecodeError:
                st.error("Erro ao interpretar resposta da IA. Tente novamente.")
            except Exception as e:
                st.error(f"Erro ao analisar: {e}")

    # --- RENDERIZAÇÃO DOS CARDS ---
    if 'dashboard_dados' in st.session_state:
        dados = st.session_state.dashboard_dados
        meta = dados.get("metadados", {})

        st.divider()
        col_m1, col_m2, col_m3 = st.columns(3)
        col_m1.caption(f"Processo: {meta.get('numero')}")
        col_m2.caption(f"Autor: {meta.get('autor')}")
        col_m3.caption(f"Réu: {meta.get('reu')}")

        tarefas = dados.get("tarefas", [])
        if not tarefas:
            st.warning("Nenhuma pendência encontrada.")

        for i, tarefa in enumerate(tarefas):
            with st.container():
                tipo = tarefa['tipo']

                cor = "#ccc"
                icon = "📌"
                titulo = "Evento"
                if tipo == 'NOMEACAO':
                    cor = "#28a745"; icon = "✅"; titulo = "Nomeação Recebida"
                elif tipo == 'QUESITOS':
                    cor = "#ffc107"; icon = "❓"; titulo = "Quesitos Apresentados"
                elif tipo == 'HONORARIOS':
                    cor = "#17a2b8"; icon = "💰"; titulo = "Proposta de Honorários"

                # Sanitiza conteúdo vindo da IA antes de renderizar como HTML
                descricao_safe = html.escape(tarefa['descricao'])
                pagina_safe = html.escape(str(tarefa['pagina']))

                st.markdown(f"""
                <div style="background-color: #262730; padding: 15px; border-radius: 8px; border-left: 6px solid {cor}; margin-bottom: 15px;">
                    <h4 style="color:white; margin:0;">{icon} {titulo} <span style="font-size:0.7em; opacity:0.8;">(Pág. {pagina_safe})</span></h4>
                    <p style="color:#ddd; margin:5px 0;">{descricao_safe}</p>
                </div>
                """, unsafe_allow_html=True)

                col_btn, col_extra = st.columns([1, 2])

                nome_perito = st.session_state.nome_perito or "Perito não informado"

                # 1. GERAR ACEITE
                if tipo == 'NOMEACAO':
                    if col_btn.button("📄 Gerar Petição de Aceite", key=f"btn_aceite_{i}"):
                        if not st.session_state.nome_perito:
                            st.warning("Informe seu nome na barra lateral antes de gerar documentos.")
                        else:
                            doc = criar_template_aceite()
                            ctx = {
                                "numero_processo": meta.get('numero'),
                                "vara": meta.get('vara'),
                                "autor": meta.get('autor'),
                                "reu": meta.get('reu'),
                                "data_atual": datetime.now().strftime("%d/%m/%Y"),
                                "nome_perito": nome_perito,
                            }
                            bio_final = renderizar_docx(doc, ctx)
                            st.download_button(
                                "⬇️ Baixar Aceite.docx",
                                bio_final.getvalue(),
                                "Aceite.docx",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                key=f"dl_aceite_{i}"
                            )

                # 2. EXTRAIR CADERNO DE QUESITOS
                elif tipo == 'QUESITOS':
                    if col_btn.button("📝 Extrair Caderno de Quesitos", key=f"btn_quesitos_{i}"):
                        doc = Document()
                        doc.add_heading("CADERNO DE QUESITOS", 0)
                        doc.add_paragraph(f"Referência: Pág. {tarefa['pagina']}")

                        lista = tarefa.get('lista_quesitos', [])
                        if isinstance(lista, list):
                            for q_item in lista:
                                p = doc.add_paragraph()
                                run = p.add_run(str(q_item))
                                run.bold = True
                                doc.add_paragraph("RESPOSTA: ___________________________________________________\n")
                        else:
                            doc.add_paragraph(str(lista))

                        bio = io.BytesIO()
                        doc.save(bio)
                        st.download_button(
                            "⬇️ Baixar Caderno de Quesitos.docx",
                            bio.getvalue(),
                            "Quesitos.docx",
                            key=f"dl_quesitos_{i}"
                        )

                # 3. PROPOSTA DE HONORÁRIOS
                elif tipo == 'HONORARIOS':
                    col_extra.caption("Calculadora Rápida:")
                    c1, c2 = col_extra.columns(2)
                    horas = c1.number_input("Total Horas", min_value=1, value=10, key=f"hs_{i}")
                    valor = c2.number_input("Valor Hora (R$)", min_value=1, value=300, key=f"vl_{i}")
                    total = horas * valor
                    c1.markdown(f"**Total: R$ {total:,.2f}**")

                    if col_btn.button("💰 Gerar Proposta", key=f"btn_hon_{i}"):
                        if not st.session_state.nome_perito:
                            st.warning("Informe seu nome na barra lateral antes de gerar documentos.")
                        else:
                            doc = criar_template_honorarios()
                            ctx = {
                                "numero_processo": meta.get('numero'),
                                "vara": meta.get('vara'),
                                "nome_perito": nome_perito,
                                "horas_vistoria": int(horas * 0.4),
                                "horas_analise": int(horas * 0.3),
                                "horas_redacao": int(horas * 0.3),
                                "total_horas": horas,
                                "valor_hora": f"{valor:,.2f}",
                                "valor_total": f"{total:,.2f}",
                            }
                            bio_final = renderizar_docx(doc, ctx)
                            st.download_button(
                                "⬇️ Baixar Proposta.docx",
                                bio_final.getvalue(),
                                "Proposta_Honorarios.docx",
                                key=f"dl_hon_{i}"
                            )


# ==============================================================================
# FERRAMENTAS MANUAIS
# ==============================================================================
if selected == "Ferramentas Manuais":
    st.subheader("🛠️ Ferramentas Avulsas")
    tab1, tab2 = st.tabs(["Calculadora de Prazos", "Extrair Quesitos (Manual)"])

    with tab1:
        data_intimacao = st.date_input("Data da Intimação")
        prazo_dias = st.number_input("Dias Úteis", min_value=1, value=15)
        if st.button("Calcular"):
            data_vencimento = calcular_prazo_uteis(
                datetime.combine(data_intimacao, datetime.min.time()),
                int(prazo_dias)
            )
            st.success(f"Vence em: {data_vencimento.strftime('%d/%m/%Y')}")
            st.caption("⚠️ Atenção: feriados nacionais e estaduais não são considerados neste cálculo.")

    with tab2:
        st.info("Esta funcionalidade está disponível no **Dashboard de Processos** com análise completa do PDF.")
