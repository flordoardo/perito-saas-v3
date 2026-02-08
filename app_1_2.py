import streamlit as st
import pdfplumber
import google.generativeai as genai
import json
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Cm
import io
import os
from datetime import datetime, timedelta
import pandas as pd
from streamlit_option_menu import option_menu # <--- BIBLIOTECA NOVA

# --- CONFIGURAÇÃO INICIAL ---
st.set_page_config(page_title="PeritoSaaS Pro", page_icon="⚖️", layout="wide")

# --- FUNÇÕES UTILITÁRIAS (Mantivemos iguais) ---
def garantir_template_padrao():
    if not os.path.exists("template_padrao.docx"):
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(3); section.bottom_margin = Cm(2)
            section.left_margin = Cm(3); section.right_margin = Cm(2)
        doc.add_paragraph('{{ cabecalho_imagem }}')
        doc.add_heading('PETIÇÃO', 0)
        doc.add_paragraph('Exmo. Juiz da {{vara}}')
        doc.add_paragraph('Proc. {{numero_processo}}')
        doc.add_paragraph('Autor: {{autor}} | Réu: {{reu}}')
        doc.add_paragraph('\n{{ corpo_peticao }}')
        doc.add_paragraph('\nBelém, {{ data_atual }}.')
        doc.add_paragraph('\n___________________________\n{{ nome_perito }}\nPerito')
        doc.save("template_padrao.docx")

def calcular_prazo_uteis(data_inicio, dias):
    dias_uteis = 0
    data_atual = data_inicio
    while dias_uteis < dias:
        data_atual += timedelta(days=1)
        if data_atual.weekday() < 5: 
            dias_uteis += 1
    return data_atual

# --- CABEÇALHO DO APP ---
# Removemos o st.title gigante para ganhar espaço
st.markdown("<h3 style='text-align: center; color: #333;'>⚖️ PeritoSaaS: Suíte Profissional</h3>", unsafe_allow_html=True)
st.markdown("---")

api_key = os.environ.get("GEMINI_API_KEY") or st.secrets.get("GEMINI_API_KEY")
if not api_key:
    st.error("⚠️ API Key não configurada.")
    st.stop()

# --- MENU DE NAVEGAÇÃO PROFISSIONAL ---
selected = option_menu(
    menu_title=None,  # Esconde o título do menu
    options=["Gerador de Aceite", "Extrator de Quesitos", "Calculadora de Prazos"], 
    icons=["file-earmark-text", "list-task", "calendar-check"], # Ícones Bootstrap
    menu_icon="cast", 
    default_index=0, 
    orientation="horizontal",
    styles={
        "container": {"padding": "0!important", "background-color": "#f8f9fa"},
        "icon": {"color": "orange", "font-size": "18px"}, 
        "nav-link": {"font-size": "16px", "text-align": "center", "margin":"0px", "--hover-color": "#eee"},
        "nav-link-selected": {"background-color": "#FF4B4B"}, # Cor vermelha do Streamlit
    }
)

# ==============================================================================
# OPÇÃO 1: GERADOR DE ACEITE
# ==============================================================================
if selected == "Gerador de Aceite":
    st.subheader("📄 Gerar Petição de Aceite")
    
    col_upload, col_modelo = st.columns([1, 1])
    with col_upload:
        uploaded_file_aceite = st.file_uploader("PDF da Nomeação", type="pdf", key="pdf_aceite")
    
    with col_modelo:
        tipo_modelo = st.radio("Modelo:", ("Padrão do Sistema", "Meu Modelo (.docx)"), horizontal=True)
        arquivo_modelo = None
        if tipo_modelo == "Meu Modelo (.docx)":
            arquivo_modelo = st.file_uploader("Seu Modelo", type="docx", key="modelo_docx")
    
    if uploaded_file_aceite and st.button("Gerar Aceite", type="primary"):
        with st.spinner("Lendo processo..."):
            try:
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('models/gemini-flash-latest')
                with pdfplumber.open(uploaded_file_aceite) as pdf:
                    texto = "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
                
                prompt = f"""
                Extraia do texto jurídico abaixo em JSON:
                {{
                    "numero_processo": "...", "autor": "...", "reu": "...", "vara": "...",
                    "texto_aceite": "Escreva um parágrafo formal de aceitação."
                }}
                Texto: {texto[:10000]}
                """
                resp = model.generate_content(prompt)
                texto_limpo = resp.text.replace("```json", "").replace("```", "").strip()
                st.session_state.dados_aceite = json.loads(texto_limpo)
                st.success("Dados extraídos!")
            except Exception as e:
                st.error(f"Erro: {e}")

    # Edição e Download
    if 'dados_aceite' in st.session_state:
        d = st.session_state.dados_aceite
        with st.expander("📝 Editar Dados", expanded=True):
            col1, col2 = st.columns(2)
            proc = col1.text_input("Processo", d.get("numero_processo"), key="p1")
            vara = col2.text_input("Vara", d.get("vara"), key="v1")
            texto = st.text_area("Texto", d.get("texto_aceite"), height=100)
        
        if st.button("💾 Baixar Aceite (.docx)"):
            garantir_template_padrao()
            template_final = arquivo_modelo if arquivo_modelo else "template_padrao.docx"
            doc = DocxTemplate(template_final)
            ctx = {"numero_processo": proc, "vara": vara, "corpo_peticao": texto, "data_atual": datetime.now().strftime("%d/%m/%Y"), "nome_perito": "Dr. Perito"}
            doc.render(ctx)
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("Download", bio.getvalue(), "Aceite.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

# ==============================================================================
# OPÇÃO 2: EXTRATOR DE QUESITOS
# ==============================================================================
if selected == "Extrator de Quesitos":
    st.subheader("❓ Extrair Quesitos das Partes")
    uploaded_file_quesitos = st.file_uploader("PDF com os Quesitos", type="pdf", key="pdf_quesitos")
    
    if uploaded_file_quesitos and st.button("🔍 Encontrar Quesitos"):
        with st.spinner("Analisando perguntas..."):
            try:
                genai.configure(api_key=api_key)
                model = genai.GenerativeModel('models/gemini-flash-latest')
                with pdfplumber.open(uploaded_file_quesitos) as pdf:
                    texto = "\n".join([p.extract_text() for p in pdf.pages if p.extract_text()])
                
                prompt = f"""
                Analise o texto e extraia os quesitos. JSON apenas:
                {{
                    "quesitos_autor": ["1. ..."],
                    "quesitos_reu": ["1. ..."],
                    "quesitos_juizo": ["1. ..."]
                }}
                Texto: {texto[:20000]}
                """
                resp = model.generate_content(prompt)
                q_dados = json.loads(resp.text.replace("```json", "").replace("```", "").strip())
                st.session_state.quesitos = q_dados
            except Exception as e:
                st.error(f"Erro: {e}")

    if 'quesitos' in st.session_state:
        q = st.session_state.quesitos
        col_q1, col_q2 = st.columns(2)
        with col_q1:
            st.markdown("**Do Autor:**")
            st.text_area("Autor", "\n".join(q.get("quesitos_autor", [])), height=200)
        with col_q2:
            st.markdown("**Do Réu:**")
            st.text_area("Réu", "\n".join(q.get("quesitos_reu", [])), height=200)
            
        if st.button("💾 Baixar Laudo com Quesitos (.docx)"):
            doc = Document()
            doc.add_heading("LAUDO PERICIAL", 0)
            doc.add_heading("1. QUESITOS DO AUTOR", level=1)
            for item in q.get("quesitos_autor", []):
                doc.add_paragraph(item)
                doc.add_paragraph("RESPOSTA: _______________________")
            doc.add_heading("2. QUESITOS DO RÉU", level=1)
            for item in q.get("quesitos_reu", []):
                doc.add_paragraph(item)
                doc.add_paragraph("RESPOSTA: _______________________")
            bio = io.BytesIO()
            doc.save(bio)
            st.download_button("Download Laudo Base", bio.getvalue(), "Laudo_Quesitos.docx")

# ==============================================================================
# OPÇÃO 3: CALCULADORA DE PRAZOS
# ==============================================================================
if selected == "Calculadora de Prazos":
    st.subheader("🗓️ Calculadora de Prazos Processuais")
    
    col_data, col_dias = st.columns(2)
    data_inicio = col_data.date_input("Data da Intimação/Leitura")
    dias_prazo = col_dias.number_input("Prazo em Dias Úteis", value=15, step=5)
    
    if st.button("Calcular Vencimento"):
        dt_inicio = datetime.combine(data_inicio, datetime.min.time())
        vencimento = calcular_prazo_uteis(dt_inicio, dias_prazo)
        st.success(f"📅 Prazo Final: **{vencimento.strftime('%d/%m/%Y')}** ({vencimento.strftime('%A')})")